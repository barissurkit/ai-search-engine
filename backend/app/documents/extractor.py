from io import BytesIO
from pathlib import PurePath

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.documents.models import DocumentExtractionError, ExtractedDocument

SUPPORTED_EXTENSIONS = {".pdf": "application/pdf", ".txt": "text/plain", ".md": "text/markdown", ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}


def extract_document(filename: str, content: bytes, max_characters: int = 500_000, max_pdf_pages: int = 100) -> ExtractedDocument:
    extension = PurePath(filename).suffix.lower()
    media_type = SUPPORTED_EXTENSIONS.get(extension)
    if media_type is None:
        raise DocumentExtractionError("Unsupported file type. Upload PDF, TXT, Markdown, or DOCX.")
    if extension == ".pdf":
        if not content.startswith(b"%PDF"):
            raise DocumentExtractionError("The PDF file is invalid.")
        reader = PdfReader(BytesIO(content))
        if len(reader.pages) > max_pdf_pages:
            raise DocumentExtractionError("PDF has too many pages.")
        pages = [(index + 1, page.extract_text() or "") for index, page in enumerate(reader.pages)]
    elif extension == ".docx":
        document = DocxDocument(BytesIO(content))
        values = [paragraph.text for paragraph in document.paragraphs]
        values.extend(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        pages = [(None, "\n".join(values))]
    else:
        pages = [(None, content.decode("utf-8", errors="replace"))]
    normalized = [(page, text.strip()) for page, text in pages if text.strip()]
    if not normalized:
        raise DocumentExtractionError("Document has no extractable text.")
    if sum(len(text) for _, text in normalized) > max_characters:
        raise DocumentExtractionError("Extracted document text is too large.")
    return ExtractedDocument(filename=PurePath(filename).name, media_type=media_type, pages=normalized)
